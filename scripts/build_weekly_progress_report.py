from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "IBM_Weekly_Progress_Report_IT_Support_Copilot.docx"
DIAGRAM_PATH = OUT_DIR / "enterprise_it_support_architecture.png"


BLUE = "1F4E79"
CYAN = "D9EAF7"
PALE = "F4F7FB"
GREEN = "DDEEDB"
GOLD = "FFF2CC"
GRAY = "666666"
BLACK = "111111"
WHITE = "FFFFFF"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_rounded(draw, box, fill, outline=BLUE, width=2, radius=16):
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def draw_centered(draw, box, text, fill=BLACK, size=28, bold=False, line_gap=4):
    fnt = font(size, bold)
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = box[1] + ((box[3] - box[1]) - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        x = box[0] + ((box[2] - box[0]) - w) / 2
        draw.text((x, y), line, font=fnt, fill=f"#{fill}")
        y += h + line_gap


def arrow(draw, start, end, fill=BLUE, width=4):
    draw.line([start, end], fill=f"#{fill}", width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - direction * 16, ey - 9), (ex - direction * 16, ey + 9)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 9, ey - direction * 16), (ex + 9, ey - direction * 16)]
    draw.polygon(points, fill=f"#{fill}")


def build_diagram():
    img = Image.new("RGB", (1800, 620), f"#{WHITE}")
    draw = ImageDraw.Draw(img)

    draw.text((52, 34), "Enterprise IT Support Copilot - Architecture", font=font(34, True), fill=f"#{BLUE}")
    draw.text(
        (52, 78),
        "Controlled RAG backend using IBM watsonx.ai, OpenSearch, FastAPI, LangGraph, and approved documentation",
        font=font(22),
        fill=f"#{GRAY}",
    )

    boxes = {
        "ui": (60, 180, 310, 300),
        "api": (400, 180, 650, 300),
        "graph": (740, 150, 1070, 330),
        "retrieval": (1180, 110, 1500, 235),
        "gen": (1180, 280, 1500, 405),
        "response": (1570, 180, 1760, 300),
        "sources": (60, 440, 335, 560),
        "ingest": (430, 440, 705, 560),
        "index": (800, 440, 1075, 560),
    }

    draw_rounded(draw, boxes["ui"], CYAN)
    draw_centered(draw, boxes["ui"], "User\nwatsonx Orchestrate", size=25, bold=True)
    draw_rounded(draw, boxes["api"], PALE)
    draw_centered(draw, boxes["api"], "FastAPI\nPOST /v1/assist", size=25, bold=True)
    draw_rounded(draw, boxes["graph"], GOLD)
    draw_centered(
        draw,
        boxes["graph"],
        "LangGraph workflow\nclassify -> scope -> retrieve\nvalidate -> generate -> safety",
        size=23,
        bold=True,
    )
    draw_rounded(draw, boxes["retrieval"], GREEN)
    draw_centered(draw, boxes["retrieval"], "OpenSearch\nBM25 + vector kNN\nRRF fusion", size=23, bold=True)
    draw_rounded(draw, boxes["gen"], GREEN)
    draw_centered(draw, boxes["gen"], "watsonx.ai\nSlate embeddings\nGranite generation", size=23, bold=True)
    draw_rounded(draw, boxes["response"], CYAN)
    draw_centered(draw, boxes["response"], "Answer\nwith citations", size=24, bold=True)
    draw_rounded(draw, boxes["sources"], PALE)
    draw_centered(draw, boxes["sources"], "Approved sources\nOCP/SNO PDFs\nOrchestrate docs\nIBM Bob docs", size=22, bold=True)
    draw_rounded(draw, boxes["ingest"], PALE)
    draw_centered(draw, boxes["ingest"], "Ingestion pipeline\nCOS/local/web load\nparse + chunk", size=23, bold=True)
    draw_rounded(draw, boxes["index"], PALE)
    draw_centered(draw, boxes["index"], "Knowledge index\nmetadata validation\n15K+ OCP chunks", size=23, bold=True)

    arrow(draw, (310, 240), (400, 240))
    arrow(draw, (650, 240), (740, 240))
    arrow(draw, (1070, 205), (1180, 172))
    arrow(draw, (1070, 275), (1180, 342))
    arrow(draw, (1500, 172), (1570, 220))
    arrow(draw, (1500, 342), (1570, 260))
    arrow(draw, (335, 500), (430, 500))
    arrow(draw, (705, 500), (800, 500))
    arrow(draw, (940, 440), (1300, 235))

    img.save(DIAGRAM_PATH)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, size=8.4, bold=False, color=BLACK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D9E2EC", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Compact Bullet")
        p.paragraph_format.space_after = Pt(1.3)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(item)
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        r.font.size = Pt(8.7)


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=5)
    set_table_borders(table)
    labels = [
        ("Domains", "OCP/SNO + Orchestrate + Bob"),
        ("Backend", "FastAPI + LangGraph"),
        ("Retrieval", "OpenSearch BM25 + kNN"),
        ("Evaluation", "38/40 passed = 95%"),
        ("Tests", "58+ focused tests passing"),
    ]
    for i, (label, value) in enumerate(labels):
        cell = table.cell(0, i)
        set_cell_shading(cell, "EEF4FB")
        set_cell_text(cell, f"{label}\n{value}", size=7.8, bold=True, color=BLACK)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9)

    if "Compact Bullet" not in styles:
        style = styles.add_style("Compact Bullet", 1)
    else:
        style = styles["Compact Bullet"]
    style.base_style = styles["Normal"]
    style.paragraph_format.left_indent = Inches(0.18)
    style.paragraph_format.first_line_indent = Inches(-0.1)
    style.paragraph_format.space_after = Pt(1.3)
    style.paragraph_format.line_spacing = 1.0


def build_docx():
    build_diagram()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Weekly Progress Report - Enterprise IT Support Copilot")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("IBM Internship POC | Status through 10 July 2026 | Vaibhav Janga")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_picture(str(DIAGRAM_PATH), width=Inches(9.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_metric_strip(doc)

    body = doc.add_table(rows=1, cols=3)
    set_table_borders(body, color="FFFFFF", size="0")
    widths = [Inches(3.15), Inches(3.35), Inches(3.15)]
    for idx, width in enumerate(widths):
        for cell in body.columns[idx].cells:
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    sections = [
        (
            "Started With",
            [
                "Built a controlled RAG backend for OpenShift Container Platform and Single Node OpenShift support.",
                "Created FastAPI service with /v1/assist, /healthz, /readyz, API-key auth, Pydantic schemas, and OpenAPI spec for watsonx Orchestrate.",
                "Designed a bounded LangGraph workflow for classification, scope resolution, retrieval, evidence checks, answer generation, citation validation, and safety handling.",
                "Implemented watsonx.ai providers for Slate embeddings and Granite generation.",
            ],
        ),
        (
            "Core Delivery",
            [
                "Built PDF ingestion from IBM Cloud Object Storage/local files using parsing, chunking, metadata validation, embeddings, and OpenSearch indexing.",
                "Implemented hybrid retrieval using BM25 keyword search, vector kNN search, and Reciprocal Rank Fusion.",
                "Indexed the approved OCP/SNO corpus, including installation, networking, storage, troubleshooting, authentication, operators, and update guides.",
                "Added deterministic refusal/clarification behavior for ambiguous, unsupported, out-of-scope, or insufficient-evidence requests.",
            ],
        ),
        (
            "Latest Additions",
            [
                "Expanded the product into a multi-domain Enterprise IT Support Copilot covering OCP/SNO, IBM watsonx Orchestrate, and IBM Bob.",
                "Added web documentation ingestion through web source discovery plus HTML, Markdown, and plain-text parsing.",
                "Added new corpus manifests for watsonx Orchestrate and IBM Bob, plus multi-domain taxonomy, classifier, and scope routing updates.",
                "Fixed non-OCP source labels and Orchestrate payload normalization so IBM Bob/Orchestrate questions route correctly and return product-aware citations.",
            ],
        ),
    ]
    for idx, (heading, bullets) in enumerate(sections):
        cell = body.cell(0, idx)
        set_cell_shading(cell, "FFFFFF")
        add_heading(cell, heading)
        add_bullets(cell, bullets)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(2)
    footer.paragraph_format.space_after = Pt(0)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        "Current status: working citation-grounded support backend with validated OpenShift/SNO evaluation, multi-domain expansion, Orchestrate integration guidance, Docker/Code Engine deployment docs, and documented next work for broader Bob/Orchestrate evaluation and cross-version retrieval."
    )
    fr.font.name = "Arial"
    fr._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    fr._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    fr.font.size = Pt(8.2)
    fr.font.color.rgb = RGBColor.from_string(GRAY)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
